# -*- coding: utf-8 -*-
"""
Created on Sun Jul  9 10:14:15 2023

@author: likun
"""
'''
This script is used to mine the core of a PDB structure, and the following three lines need to be modified before running.
line98 corefile='E:/imm/inf_core_sel2.txt' #Pathogen epitope cores(PECs), which can be found in https://github.com/LKbiosoup/epitope-core obtain
line99 pdb='E:/pdb2/spike_rbd_wt.pdb'  # A PDB structure that needs to be input, such as the RBD region of spike protein 
line100 min_RSA=0.1  #Filter for epitope cores with RSA greater than this value
'''
from docx import Document
from docx.shared import RGBColor
from Bio.PDB import PDBParser
from Bio.PDB import SASA
import Bio.PDB
import docx
from docx.shared import RGBColor
import os

def extract_red_strings(filename):
    doc = docx.Document(filename)
    red_strings = []
    current_red_string = ""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color.rgb == RGBColor(255, 0, 0):
                current_red_string += run.text
            else:
                if current_red_string:
                    red_strings.append(current_red_string)
                    current_red_string = ""
        if current_red_string:
            red_strings.append(current_red_string)
            current_red_string = ""
    return red_strings


def three_one(string):
    aa_3code = ['ALA', 'ARG', 'ASN', 'ASP', 'CYS', 'GLU', 'GLN', 'GLY', 
            'HIS', 'ILE', 'LEU', 'LYS', 'MET', 'PHE', 'PRO', 'SER', 
            'THR', 'TRP', 'TYR', 'VAL','HOH']
    aa_1code = ['A', 'R', 'N', 'D', 'C', 'E', 'Q', 'G', 
            'H', 'I', 'L', 'K', 'M', 'F', 'P', 'S', 
            'T', 'W', 'Y', 'V','']
    try:
        ind=aa_3code.index(string)
        return aa_1code[ind]
    except:
        return ''

def RSA_mover(pdb):
    seq=''
    sasa_lis=[]
    parser = PDBParser()
    structure = parser.get_structure("1LCD", pdb)
    sr = SASA.ShrakeRupley()
    sr.compute(structure, level="R")
    
    for residue in structure.get_residues():
        resname = residue.get_resname()
        resname=three_one(resname)
        sasa=residue.sasa
        sasa_lis.append(sasa)
        seq=seq+resname
    sasa_lis=[i/max(sasa_lis) for i in sasa_lis]
    return seq,sasa_lis
def RSA_ave(stringlis,seq,sasa_lis):
    out=[]
    for string in stringlis:
        try:
            ind=seq.index(string)
            sublis=sasa_lis[ind:ind+len(string)]
            ave=sum(sublis)/len(sublis)
            out.append([string,ave])
        except:
            error='core not find'
    return out

def cl(seq):
    seq=seq.replace('"','')
    seq=seq.split('+')[0]
    seq=seq.strip()
    return seq


# 
def write_doc(str, list):
    p=doc.add_paragraph()
    for s in range(len(str)):
        if s in list:
            run=p.add_run(str[s])         
            run.font.color.rgb = RGBColor(255,0,0)
        else:
            run=p.add_run(str[s])
    
corefile='E:/imm/inf_core_sel2.txt'
pdb='E:/pdb2/spike_rbd_wt.pdb'  #
min_RSA=0.1  #Filter for epitope cores with RSA greater than this value
try:
    ind=corefile.rfind('/')# 一个需要指定目录临时的运行文件
    docx_file=corefile[0:ind]+'/'+'core.docx'
except:
    print('The working directory should be a secondary directory!')

seq,rsa=RSA_mover(pdb)


cores=[]
infs=[]
with open(corefile,'r') as f:
    lines=f.readlines()
f.close()
for line in lines:
    cores.append(line.strip())

doc=Document()


indx=[]
corex=[t for t in cores if t in seq]
for m in corex:
    #
    count = seq.count(m)
    inds = [seq.index(m, i) for i in range(count)] 
    for ind in inds:
        rangex=[i for i in range(ind,ind+len(m))]
        indx.extend(rangex)
indx=list(set(indx))
write_doc(seq,indx)


doc.save(docx_file)


strings=extract_red_strings(docx_file)
os.remove(docx_file)
print("All antigenic epitope cores")
lis=RSA_ave(strings, seq, rsa)
lis=[i for i in lis if i[1]>min_RSA]
for i in lis:
    print('core:',i[0],'; RSA:',i[1])

